"""Design document services."""
from __future__ import annotations
import io
import json
import logging
from pathlib import Path
from autopilot_platform.platform.core.settings import design_uploads_root
logger = logging.getLogger(__name__)
ALLOWED_DOC_EXT = {'.txt', '.md', '.csv', '.json', '.yaml', '.yml', '.docx', '.pdf', '.xlsx', '.xls'}


def uploads_root() -> Path:
    return design_uploads_root()

def extract_text_from_bytes(filename: str, data: bytes) -> str:
    ext = Path(filename or '').suffix.lower()
    if ext not in ALLOWED_DOC_EXT:
        raise ValueError(f"不支持的文件类型: {ext or '(无扩展名)'}")
    if ext in {'.txt', '.md', '.csv', '.yaml', '.yml'}:
        return data.decode('utf-8', errors='ignore')
    if ext == '.json':
        try:
            obj = json.loads(data.decode('utf-8', errors='ignore'))
            return json.dumps(obj, ensure_ascii=False, indent=2)
        except json.JSONDecodeError:
            return data.decode('utf-8', errors='ignore')
    if ext == '.docx':
        try:
            from docx import Document  # 延迟：可选 extra
            doc = Document(io.BytesIO(data))
            docx_parts: list[str] = []
            for p in doc.paragraphs:
                t = (p.text or '').strip()
                if t:
                    docx_parts.append(t)
            for table in doc.tables:
                for row in table.rows:
                    cells = [(c.text or '').strip() for c in row.cells]
                    cells = [c for c in cells if c]
                    if cells:
                        docx_parts.append(' | '.join(cells))
            return '\n'.join(docx_parts)
        except ImportError as exc:
            raise ValueError('解析 docx 需要安装 python-docx') from exc
    if ext == '.pdf':
        try:
            from PyPDF2 import PdfReader  # 延迟：可选 extra
            reader = PdfReader(io.BytesIO(data))
            pdf_parts: list[str] = []
            for page in reader.pages:
                t = (page.extract_text() or '').strip()
                if t:
                    pdf_parts.append(t)
            return '\n'.join(pdf_parts)
        except ImportError as exc:
            raise ValueError('解析 pdf 需要安装 PyPDF2') from exc
    if ext in {'.xlsx', '.xls'}:
        try:
            from openpyxl import load_workbook  # 延迟：可选 extra
            wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
            xls_parts: list[str] = []
            for sheet in wb.worksheets:
                xls_parts.append(f'# {sheet.title}')
                for row in sheet.iter_rows(values_only=True):
                    cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                    if cells:
                        xls_parts.append(' | '.join(cells))
            return '\n'.join(xls_parts)
        except ImportError as exc:
            raise ValueError('解析 xlsx 需要安装 openpyxl') from exc
        except Exception as exc:
            raise ValueError(f'Excel 无法解析: {exc}') from exc
    return data.decode('utf-8', errors='ignore')
