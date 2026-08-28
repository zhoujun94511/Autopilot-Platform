#!/usr/bin/python
# -*- coding: utf-8 -*-
import json
import yaml
import pandas as pd
import openpyxl
from docx import Document
from pathlib import Path

try:
    from PyPDF2 import PdfReader

    PDF_SUPPORT = True
except ImportError:
    PdfReader = None  # 确保名字总是存在
    PDF_SUPPORT = False


def file_loader(file_path):
    """
    通用文件加载器，支持 txt / md / csv / docx / json / xlsx / pdf / yaml
    返回: text_list(纯文本列表), doc_splits(Document分块), metadata_list(元数据列表)
    """
    ext = Path(file_path).suffix.lower()

    try:
        # ----------------- 文本 / Markdown -----------------
        if ext in [".txt", ".md"]:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                lines = [line.strip() for line in f.readlines() if line.strip()]
            text_chunks = lines
            file_metadata = {"source": file_path, "type": ext[1:]}

        # ----------------- CSV -----------------
        elif ext == ".csv":
            df = pd.read_csv(file_path)
            rows = [" ".join(map(str, row)) for row in df.values]
            text_chunks = rows
            file_metadata = {"source": file_path, "type": "csv", "rows": len(rows)}

        # ----------------- Word -----------------
        elif ext == ".docx":
            from utils.utils_core.common_utils import clean_text
            doc = Document(file_path)
            text_chunks = []
            
            # 1. 提取段落文本
            for p in doc.paragraphs:
                text = p.text.strip()
                if text:
                    # 清理乱码字符
                    cleaned_text = clean_text(text)
                    if cleaned_text.strip():
                        text_chunks.append(cleaned_text)
            
            # 2. 提取表格文本
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_cells = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            # 清理乱码字符
                            cleaned_cell = clean_text(cell_text)
                            if cleaned_cell.strip():
                                row_cells.append(cleaned_cell)
                    if row_cells:
                        # 将表格行转换为文本，用制表符分隔
                        table_rows.append(" | ".join(row_cells))
                if table_rows:
                    # 将表格作为整体文本块添加
                    text_chunks.append("\n".join(table_rows))
            
            file_metadata = {"source": file_path, "type": "docx"}

        # ----------------- JSON -----------------
        elif ext == ".json":
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                data = json.load(f)
            # 拆成 key:value 形式
            lines = []

            def flatten_json(d, prefix=""):
                if isinstance(d, dict):
                    for k, v in d.items():
                        flatten_json(v, f"{prefix}{k}.")
                elif isinstance(d, list):
                    for idx, v in enumerate(d):
                        flatten_json(v, f"{prefix}{idx}.")
                else:
                    lines.append(f"{prefix[:-1]}: {d}")

            flatten_json(data)
            text_chunks = lines if lines else [json.dumps(data, ensure_ascii=False, indent=2)]
            file_metadata = {"source": file_path, "type": "json"}
        # ----------------- Excel -----------------
        elif ext == ".xlsx":
            wb = openpyxl.load_workbook(file_path, read_only=True)
            rows = []
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                for row in ws.iter_rows(values_only=True):
                    row_text = " ".join([str(cell) for cell in row if cell is not None])
                    if row_text.strip():
                        rows.append(row_text)
            text_chunks = rows
            file_metadata = {"source": file_path, "type": "xlsx", "rows": len(rows)}

        # ----------------- PDF -----------------
        elif ext == ".pdf":
            if not PDF_SUPPORT:
                raise ImportError("PyPDF2 未安装，无法解析 PDF")
            from utils.utils_core.common_utils import clean_text
            reader = PdfReader(file_path)
            rows = []
            for page_num, page in enumerate(reader.pages, start=1):
                text = page.extract_text() or ""
                # 清理PDF提取的文本，移除乱码字符
                text = clean_text(text)
                if text.strip():
                    page_lines = [line.strip() for line in text.splitlines() if line.strip()]
                    rows.extend(page_lines)
            text_chunks = rows
            file_metadata = {"source": file_path, "type": "pdf", "pages": len(reader.pages)}

        # ----------------- YAML -----------------
        elif ext in [".yaml", ".yml"]:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                data = yaml.safe_load(f)
            text_chunks = [f"{k}: {v}" for k, v in data.items()] if isinstance(data, dict) else [
                json.dumps(data, ensure_ascii=False)]
            file_metadata = {"source": file_path, "type": "yaml"}

        else:
            raise ValueError(f"Unsupported file format: {file_path}")

    except Exception as e:
        raise ValueError(f"Error loading file {file_path}: {str(e)}")

    # ✅ 分块逻辑（优化版：减少碎片化）
    # 尝试导入 Document（兼容不同版本的 langchain）
    try:
        from langchain_core.documents import Document as LangchainDocument
    except ImportError:
        try:
            from langchain.schema import Document as LangchainDocument
        except ImportError:
            try:
                from langchain.documents import Document as LangchainDocument
            except ImportError:
                # 如果都失败，定义一个简单的 Document 类作为 fallback
                from typing import Dict, Any, Optional
                class LangchainDocument:
                    """Document 类（fallback）"""
                    def __init__(self, page_content: str, metadata: Optional[Dict[str, Any]] = None):
                        # 参数名 metadata 与外部作用域变量同名，但这是 LangChain Document 的标准接口
                        # 必须保持参数名为 metadata 以保持接口兼容性
                        # 立即将参数赋值给局部变量，避免变量名隐藏警告
                        doc_metadata = metadata or {} if metadata is not None else {}
                        self.page_content = page_content
                        self.metadata = doc_metadata

    doc_splits = []
    # 增大块大小减少碎片化，增大重叠避免重要内容被切断
    chunk_size = 1000  # 原来500，现在1000
    overlap = 200      # 原来100，现在200

    for i, text in enumerate(text_chunks):
        if len(text) <= chunk_size:
            doc_splits.append(LangchainDocument(
                page_content=text,
                metadata=file_metadata
            ))
        else:
            start = 0
            while start < len(text):
                end = min(start + chunk_size, len(text))
                piece = text[start:end]
                doc_splits.append(LangchainDocument(
                    page_content=piece,
                    metadata=file_metadata
                ))
                start = end - overlap if end < len(text) else end

    text_list = [doc.page_content for doc in doc_splits]
    metadata_list = [doc.metadata for doc in doc_splits]

    return text_list, doc_splits, metadata_list
